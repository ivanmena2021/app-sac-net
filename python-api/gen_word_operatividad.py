"""
gen_word_operatividad.py — Generador de Ayuda Memoria OPERATIVIDAD SAC
======================================================================
Versión 2.0 — Con gráficos embebidos (matplotlib), numeración correlativa,
diseño profesional y contenido completo fiel al formato oficial.

Incluye:
  - Gráfico 1: Avisos por departamento (barras)
  - Gráfico 2: Avisos por tipo de siniestro (barras)
  - Gráfico 3: Desembolsos vs Indemnización por empresa (barras agrupadas)
  - Cuadros 1-5 con formato profesional
  - Texto narrativo completo
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker


# ═══ PALETA DE COLORES ═══
AZUL_OSCURO = "#1F4E79"
AZUL_MEDIO = "#2E75B6"
AZUL_CLARO = "#D6E4F0"
AMARILLO = "#FFC000"
GRIS = "#666666"
BLANCO = "#FFFFFF"
NEGRO = "#000000"

C = {
    "HEADER_BG": "1F4E79",
    "SUBTOTAL_BG": "D6E4F0",
    "ALT_ROW": "F2F7FB",
    "WHITE": "FFFFFF",
    "BLACK": "000000",
}


# ═══════════════════════════════════════════════════════════════════
# FORMATO DE NÚMEROS
# ═══════════════════════════════════════════════════════════════════

def fmt(val, dec=2):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "S/ -"
    try:
        n = float(val)
        if n == 0:
            return "S/ -"
        return f"S/ {n:,.{dec}f}"
    except (ValueError, TypeError):
        return "S/ -"


def fmt_n(val, dec=2):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "-"
    try:
        n = float(val)
        if n == 0:
            return "-"
        return f"{n:,.{dec}f}"
    except (ValueError, TypeError):
        return "-"


def fmt_pct(val, dec=2):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "0.00%"
    try:
        return f"{float(val):,.{dec}f}%"
    except (ValueError, TypeError):
        return "0.00%"


def fmt_int(val):
    try:
        n = int(float(val))
        return f"{n:,}"
    except (ValueError, TypeError):
        return "-"


# ═══════════════════════════════════════════════════════════════════
# GENERACIÓN DE GRÁFICOS (matplotlib → BytesIO PNG)
# ═══════════════════════════════════════════════════════════════════

def _chart_avisos_departamento(avisos_by_depto):
    """Gráfico de barras: Número de Avisos Reportados por Departamento."""
    fig, ax = plt.subplots(figsize=(10, 4.5))

    deptos = [d.title() for d in avisos_by_depto.index]
    valores = avisos_by_depto.values.astype(int)

    bars = ax.bar(range(len(deptos)), valores, color=AZUL_MEDIO, width=0.7, edgecolor="none")

    # Etiquetas encima de cada barra
    for bar, val in zip(bars, valores):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(valores) * 0.01,
                f"{val:,}", ha="center", va="bottom", fontsize=7, fontweight="bold",
                color=AZUL_OSCURO)

    ax.set_xticks(range(len(deptos)))
    ax.set_xticklabels(deptos, rotation=45, ha="right", fontsize=7)
    ax.set_title("Número de Avisos Reportados por Departamento",
                 fontsize=12, fontweight="bold", color=AZUL_OSCURO, pad=12)
    ax.set_ylabel("")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{int(x):,}"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CCCCCC")
    ax.spines["bottom"].set_color("#CCCCCC")
    ax.tick_params(axis="y", labelsize=7, colors="#888888")
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_axisbelow(True)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_avisos_tipo(avisos_by_tipo):
    """Gráfico de barras: Número de Avisos Reportados por Tipo de Siniestro."""
    fig, ax = plt.subplots(figsize=(10, 4.5))

    tipos = [t.title() for t in avisos_by_tipo.index]
    valores = avisos_by_tipo.values.astype(int)

    bars = ax.bar(range(len(tipos)), valores, color=AZUL_MEDIO, width=0.7, edgecolor="none")

    for bar, val in zip(bars, valores):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(valores) * 0.01,
                f"{val:,}", ha="center", va="bottom", fontsize=7, fontweight="bold",
                color=AZUL_OSCURO)

    ax.set_xticks(range(len(tipos)))
    ax.set_xticklabels(tipos, rotation=45, ha="right", fontsize=7)
    ax.set_title("Número de Avisos Reportados por Tipo de Siniestro",
                 fontsize=12, fontweight="bold", color=AZUL_OSCURO, pad=12)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{int(x):,}"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CCCCCC")
    ax.spines["bottom"].set_color("#CCCCCC")
    ax.tick_params(axis="y", labelsize=7, colors="#888888")
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_axisbelow(True)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_desembolsos_empresa(indemn_lp, desemb_lp, indemn_rimac, desemb_rimac):
    """Gráfico de barras agrupadas: Desembolsos vs Indemnización por Empresa."""
    fig, ax = plt.subplots(figsize=(8, 4))

    empresas = ["La Positiva", "Rímac"]
    indemnizaciones = [indemn_lp, indemn_rimac]
    desembolsos = [desemb_lp, desemb_rimac]

    x = np.arange(len(empresas))
    width = 0.3

    bars1 = ax.bar(x - width / 2, indemnizaciones, width, label="Suma de INDEMNIZACIÓN",
                   color=AZUL_MEDIO, edgecolor="none")
    bars2 = ax.bar(x + width / 2, desembolsos, width, label="Suma de MONTO DESEMBOLSADO",
                   color=AMARILLO, edgecolor="none")

    # Etiquetas
    for bar, val in zip(bars1, indemnizaciones):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(indemnizaciones) * 0.02,
                f"S/ {val:,.2f}", ha="center", va="bottom", fontsize=8, fontweight="bold",
                color=AZUL_OSCURO)
    for bar, val in zip(bars2, desembolsos):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(indemnizaciones) * 0.02,
                f"S/ {val:,.2f}", ha="center", va="bottom", fontsize=8, fontweight="bold",
                color="#996600")

    ax.set_xticks(x)
    ax.set_xticklabels(empresas, fontsize=10, fontweight="bold")
    ax.set_title("DESEMBOLSOS REALIZADOS A NIVEL DE EMPRESA DE SEGUROS",
                 fontsize=11, fontweight="bold", color=AZUL_OSCURO, pad=12)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.08), ncol=2, fontsize=8,
              frameon=True, fancybox=True)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"S/ {x:,.0f}"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CCCCCC")
    ax.spines["bottom"].set_color("#CCCCCC")
    ax.tick_params(axis="y", labelsize=7, colors="#888888")
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.set_axisbelow(True)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════
# UTILIDADES PARA TABLAS WORD
# ═══════════════════════════════════════════════════════════════════

def _set_bg(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_borders(cell, color="AAAAAA"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    tcPr.append(tcW)


def _write_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
                font_color=None, bg_color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.name = "Arial Narrow"
    run.font.size = Pt(size)
    run.font.bold = bold
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)
    if bg_color:
        _set_bg(cell, bg_color)
    _set_borders(cell)


# ═══════════════════════════════════════════════════════════════════
# UTILIDADES PARA TEXTO WORD
# ═══════════════════════════════════════════════════════════════════

def _add_section_heading(doc, text):
    """Título de sección con línea inferior azul."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor.from_string("000000")
    return p


def _add_bullet(doc, text, bold_start=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if bold_start:
        run = p.add_run(bold_start)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def _add_subbullet(doc, text):
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def _add_body(doc, text, bold=False, italic=False, size=10, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def _add_bold_underline_bullet(doc, text):
    """Bullet con texto negrita y subrayado (para indemnizaciones)."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.underline = True
    return p


def _add_cuadro_title(doc, numero, titulo):
    """Agrega título de cuadro numerado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"Cuadro N° {numero:02d}: {titulo}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    return p


def _add_grafico_caption(doc, numero, titulo):
    """Agrega pie de gráfico numerado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Gráfico N° {numero:02d}: {titulo}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("666666")
    return p


def _add_image_centered(doc, image_buf, width_inches=6.0):
    """Agrega una imagen centrada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_buf, width=Inches(width_inches))
    return p


# ═══════════════════════════════════════════════════════════════════
# PREPARACIÓN DE DATOS
# ═══════════════════════════════════════════════════════════════════

def _prepare_operatividad_data(datos):
    midagri = datos["midagri"].copy()
    materia = datos["materia"]

    # Mapeo departamento → empresa
    depto_empresa = {}
    if "EMPRESA_ASEGURADORA" in materia.columns and "DEPARTAMENTO" in materia.columns:
        for _, row in materia.iterrows():
            d = str(row["DEPARTAMENTO"]).strip().upper()
            e = str(row["EMPRESA_ASEGURADORA"]).strip().upper()
            depto_empresa[d] = e

    if "DEPARTAMENTO" in midagri.columns:
        midagri["EMPRESA"] = midagri["DEPARTAMENTO"].map(depto_empresa).fillna("OTROS")
    else:
        midagri["EMPRESA"] = "OTROS"

    def _norm_emp(e):
        eu = str(e).upper()
        if "POSITIVA" in eu:
            return "LA POSITIVA"
        elif "RIMAC" in eu or "RÍMAC" in eu:
            return "RÍMAC"
        return eu
    midagri["EMPRESA"] = midagri["EMPRESA"].apply(_norm_emp)

    # Prima neta por departamento
    prima_por_depto = {}
    if "PRIMA_NETA" in materia.columns and "DEPARTAMENTO" in materia.columns:
        for _, row in materia.iterrows():
            d = str(row["DEPARTAMENTO"]).strip().upper()
            prima_por_depto[d] = float(row.get("PRIMA_NETA", 0) or 0)

    fecha_corte = datos["fecha_corte"]
    total_avisos = len(midagri)

    # Avisos por empresa
    avisos_by_empresa = midagri.groupby("EMPRESA").size().to_dict()
    avisos_lp = avisos_by_empresa.get("LA POSITIVA", 0)
    avisos_rimac = avisos_by_empresa.get("RÍMAC", 0)

    # Avisos por departamento (todos)
    avisos_by_depto = midagri.groupby("DEPARTAMENTO").size().sort_values(ascending=False)

    # Top 4 departamentos
    top4_deptos = avisos_by_depto.head(4)
    top4_total = top4_deptos.sum()
    top4_pct = (top4_total / total_avisos * 100) if total_avisos > 0 else 0

    # Avisos por tipo de siniestro (todos)
    avisos_by_tipo = midagri["TIPO_SINIESTRO"].value_counts() if "TIPO_SINIESTRO" in midagri.columns else pd.Series()

    # Top 3 tipos
    top3_tipos = avisos_by_tipo.head(3)
    top3_total = top3_tipos.sum()
    top3_pct = (top3_total / total_avisos * 100) if total_avisos > 0 else 0

    # Ajustados
    if "ESTADO_INSPECCION" in midagri.columns:
        ajustados_mask = midagri["ESTADO_INSPECCION"].astype(str).str.upper() == "CERRADO"
    elif "ESTADO_SINIESTRO" in midagri.columns:
        ajustados_mask = midagri["ESTADO_SINIESTRO"].astype(str).str.upper() == "CONCRETADO"
    else:
        ajustados_mask = pd.Series([False] * len(midagri))

    total_ajustados = int(ajustados_mask.sum())
    pct_ajustados = (total_ajustados / total_avisos * 100) if total_avisos > 0 else 0

    ajust_lp = int(ajustados_mask[midagri["EMPRESA"] == "LA POSITIVA"].sum())
    ajust_rimac = int(ajustados_mask[midagri["EMPRESA"] == "RÍMAC"].sum())
    pct_ajust_lp = (ajust_lp / avisos_lp * 100) if avisos_lp > 0 else 0
    pct_ajust_rimac = (ajust_rimac / avisos_rimac * 100) if avisos_rimac > 0 else 0

    # Indemnización total
    monto_indemnizado = midagri["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0
    sup_ind_total = midagri["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in midagri.columns else 0

    # Siniestralidad por empresa
    indemn_lp = midagri[midagri["EMPRESA"] == "LA POSITIVA"]["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0
    indemn_rimac = midagri[midagri["EMPRESA"] == "RÍMAC"]["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0
    prima_lp = sum(prima_por_depto.get(d, 0) for d in midagri[midagri["EMPRESA"] == "LA POSITIVA"]["DEPARTAMENTO"].unique())
    prima_rimac = sum(prima_por_depto.get(d, 0) for d in midagri[midagri["EMPRESA"] == "RÍMAC"]["DEPARTAMENTO"].unique())
    idx_lp = (indemn_lp / prima_lp * 100) if prima_lp > 0 else 0
    idx_rimac = (indemn_rimac / prima_rimac * 100) if prima_rimac > 0 else 0
    prima_total_neta = datos.get("prima_neta", 0)
    idx_general = (monto_indemnizado / prima_total_neta * 100) if prima_total_neta > 0 else 0

    # Desembolsos por empresa
    desemb_lp = midagri[midagri["EMPRESA"] == "LA POSITIVA"]["MONTO_DESEMBOLSADO"].sum() if "MONTO_DESEMBOLSADO" in midagri.columns else 0
    desemb_rimac = midagri[midagri["EMPRESA"] == "RÍMAC"]["MONTO_DESEMBOLSADO"].sum() if "MONTO_DESEMBOLSADO" in midagri.columns else 0

    # ─── TABLA SINIESTRALIDAD ───
    tabla_siniestralidad = []
    for empresa in ["LA POSITIVA", "RÍMAC"]:
        df_emp = midagri[midagri["EMPRESA"] == empresa]
        deptos = sorted(df_emp["DEPARTAMENTO"].unique())
        emp_indemn, emp_sup, emp_prima = 0, 0, 0
        filas = []
        for depto in deptos:
            df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
            indemn = df_d["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
            sup = df_d["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_d.columns else 0
            prima = prima_por_depto.get(depto, 0)
            idx = (indemn / prima * 100) if prima > 0 else 0
            emp_indemn += indemn
            emp_sup += sup
            emp_prima += prima
            filas.append({"empresa": empresa, "departamento": depto.title(),
                          "indemnizacion": indemn, "sup_indemnizada": sup,
                          "prima_neta": prima, "indice": idx})
        filas.sort(key=lambda x: x["indice"], reverse=True)
        tabla_siniestralidad.extend(filas)
        idx_emp = (emp_indemn / emp_prima * 100) if emp_prima > 0 else 0
        tabla_siniestralidad.append({"empresa": f"Total {empresa}", "departamento": "",
                                     "indemnizacion": emp_indemn, "sup_indemnizada": emp_sup,
                                     "prima_neta": emp_prima, "indice": idx_emp, "is_subtotal": True})

    tabla_siniestralidad.append({"empresa": "Total general", "departamento": "",
                                 "indemnizacion": monto_indemnizado, "sup_indemnizada": sup_ind_total,
                                 "prima_neta": prima_total_neta, "indice": idx_general, "is_total": True})

    # ─── TABLA COBERTURAS ───
    tabla_coberturas = []
    if "TIPO_COBERTURA" in midagri.columns:
        for empresa in ["LA POSITIVA", "RÍMAC"]:
            df_emp = midagri[midagri["EMPRESA"] == empresa]
            deptos = sorted(df_emp["DEPARTAMENTO"].unique())
            emp_comp, emp_cat, emp_total = 0, 0, 0
            for depto in deptos:
                df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
                comp = df_d.loc[df_d["TIPO_COBERTURA"].astype(str).str.upper().str.contains("COMPLEMENT", na=False), "INDEMNIZACION"].sum()
                cat = df_d.loc[df_d["TIPO_COBERTURA"].astype(str).str.upper().str.contains("CATASTR", na=False), "INDEMNIZACION"].sum()
                tot = df_d["INDEMNIZACION"].sum()
                emp_comp += comp; emp_cat += cat; emp_total += tot
                if tot > 0:
                    tabla_coberturas.append({"empresa": empresa, "departamento": depto.title(),
                                             "complementaria": comp, "catastrofica": cat, "total": tot})
            tabla_coberturas.append({"empresa": f"Total {empresa}", "departamento": "",
                                     "complementaria": emp_comp, "catastrofica": emp_cat,
                                     "total": emp_total, "is_subtotal": True})
        tc_no_sub = [r for r in tabla_coberturas if not r.get("is_subtotal")]
        total_comp = sum(r["complementaria"] for r in tc_no_sub)
        total_cat = sum(r["catastrofica"] for r in tc_no_sub)
        tabla_coberturas.append({"empresa": "Total general", "departamento": "",
                                 "complementaria": total_comp, "catastrofica": total_cat,
                                 "total": total_comp + total_cat, "is_total": True})

    # ─── TABLA PRIORIZADOS ───
    tabla_priorizados = []
    total_prio_ind, total_noprio_ind = 0, 0
    total_prio_sup, total_noprio_sup = 0, 0
    if "PRIORIZADO" in midagri.columns:
        for empresa in ["LA POSITIVA", "RÍMAC"]:
            df_emp = midagri[midagri["EMPRESA"] == empresa]
            no_mask = df_emp["PRIORIZADO"].astype(str).str.upper().str.contains("NO", na=False)
            si_mask = ~no_mask

            for label, mask in [("PRIORIZADO", si_mask), ("NO PRIORIZADO", no_mask)]:
                sup = df_emp.loc[mask, "SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_emp.columns else 0
                ind = df_emp.loc[mask, "INDEMNIZACION"].sum() if "INDEMNIZACION" in df_emp.columns else 0
                tabla_priorizados.append({"empresa": empresa, "cultivo": label,
                                          "sup_indemnizada": sup, "indemnizacion": ind})
                if label == "PRIORIZADO":
                    total_prio_ind += ind; total_prio_sup += sup
                else:
                    total_noprio_ind += ind; total_noprio_sup += sup

            emp_sup = df_emp["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_emp.columns else 0
            emp_ind = df_emp["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_emp.columns else 0
            tabla_priorizados.append({"empresa": f"Total {empresa}", "cultivo": "",
                                      "sup_indemnizada": emp_sup, "indemnizacion": emp_ind,
                                      "is_subtotal": True})
        tabla_priorizados.append({"empresa": "Total general", "cultivo": "",
                                  "sup_indemnizada": sup_ind_total, "indemnizacion": monto_indemnizado,
                                  "is_total": True})

    # ─── TABLA DESEMBOLSOS ───
    tabla_desembolsos = []
    for empresa in ["LA POSITIVA", "RÍMAC"]:
        df_emp = midagri[midagri["EMPRESA"] == empresa]
        deptos = sorted(df_emp["DEPARTAMENTO"].unique())
        emp_indemn, emp_desemb, emp_prod = 0, 0, 0
        filas = []
        for depto in deptos:
            df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
            indemn = df_d["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
            desemb = df_d["MONTO_DESEMBOLSADO"].sum() if "MONTO_DESEMBOLSADO" in df_d.columns else 0
            prod = df_d["N_PRODUCTORES"].sum() if "N_PRODUCTORES" in df_d.columns else 0
            pct = (desemb / indemn * 100) if indemn > 0 else 0
            emp_indemn += indemn; emp_desemb += desemb; emp_prod += prod
            filas.append({"empresa": empresa, "departamento": depto.title(),
                          "indemnizacion": indemn, "desembolso": desemb,
                          "pct_desembolso": pct, "productores": int(prod)})
        filas.sort(key=lambda x: x["pct_desembolso"], reverse=True)
        tabla_desembolsos.extend(filas)
        pct_emp = (emp_desemb / emp_indemn * 100) if emp_indemn > 0 else 0
        tabla_desembolsos.append({"empresa": f"Total {empresa}", "departamento": "",
                                  "indemnizacion": emp_indemn, "desembolso": emp_desemb,
                                  "pct_desembolso": pct_emp, "productores": int(emp_prod),
                                  "is_subtotal": True})

    monto_desembolsado = datos.get("monto_desembolsado", 0)
    productores = datos.get("productores_desembolso", 0)
    pct_desembolso = datos.get("pct_desembolso", 0)
    tabla_desembolsos.append({"empresa": "Total general", "departamento": "",
                              "indemnizacion": monto_indemnizado, "desembolso": monto_desembolsado,
                              "pct_desembolso": float(pct_desembolso), "productores": int(productores),
                              "is_total": True})

    return {
        "fecha_corte": fecha_corte, "total_avisos": total_avisos,
        "avisos_lp": avisos_lp, "avisos_rimac": avisos_rimac,
        "avisos_by_depto": avisos_by_depto,
        "top4_deptos": top4_deptos, "top4_total": top4_total, "top4_pct": top4_pct,
        "avisos_by_tipo": avisos_by_tipo,
        "top3_tipos": top3_tipos, "top3_total": top3_total, "top3_pct": top3_pct,
        "total_ajustados": total_ajustados, "pct_ajustados": pct_ajustados,
        "ajust_lp": ajust_lp, "ajust_rimac": ajust_rimac,
        "pct_ajust_lp": pct_ajust_lp, "pct_ajust_rimac": pct_ajust_rimac,
        "monto_indemnizado": monto_indemnizado, "sup_ind_total": sup_ind_total,
        "idx_general": idx_general, "idx_lp": idx_lp, "idx_rimac": idx_rimac,
        "indemn_lp": indemn_lp, "indemn_rimac": indemn_rimac,
        "desemb_lp": desemb_lp, "desemb_rimac": desemb_rimac,
        "tabla_siniestralidad": tabla_siniestralidad,
        "tabla_coberturas": tabla_coberturas,
        "tabla_priorizados": tabla_priorizados,
        "total_prio_ind": total_prio_ind, "total_noprio_ind": total_noprio_ind,
        "tabla_desembolsos": tabla_desembolsos,
        "monto_desembolsado": monto_desembolsado,
        "productores": productores, "pct_desembolso": pct_desembolso,
        "deptos_con_desembolso": datos.get("deptos_con_desembolso", 0),
    }


# ═══════════════════════════════════════════════════════════════════
# FUNCIÓN AUXILIAR PARA CREAR TABLAS PROFESIONALES
# ═══════════════════════════════════════════════════════════════════

def _build_table(doc, headers, data_rows, col_widths):
    """Crea tabla profesional con header, filas alternas, subtotales y total."""
    n_rows = len(data_rows) + 1
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _write_cell(cell, h, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER,
                    font_color=C["WHITE"], bg_color=C["HEADER_BG"])
        _set_cell_width(cell, col_widths[i])

    return table


# ═══════════════════════════════════════════════════════════════════
# GENERADOR PRINCIPAL
# ═══════════════════════════════════════════════════════════════════

def generate_operatividad_docx(datos):
    d = _prepare_operatividad_data(datos)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    cuadro_num = 0
    grafico_num = 0

    # ═══════════════════════════════════════════════════════════════
    # TÍTULO
    # ═══════════════════════════════════════════════════════════════
    for text in ["AYUDA MEMORIA OPERATIVIDAD SAC",
                 "CAMPAÑA AGRÍCOLA 2025-2026",
                 f"(AL {d['fecha_corte']})"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.bold = True
        run.underline = True

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════════════════════════
    # 1. OPERATIVIDAD
    # ═══════════════════════════════════════════════════════════════
    _add_body(doc, "1.  Sobre la operatividad de las empresas de seguros en la campaña 2025-2026:", size=10)

    # ─── a) Avisos de siniestros ───
    _add_section_heading(doc, "a) Avisos de siniestros")

    _add_bullet(doc,
        f"Al {d['fecha_corte']}, se registran {fmt_int(d['total_avisos'])} avisos de siniestros "
        f"reportados por las DRA's/GRA's, que fueron recepcionados por las dos (2) empresas "
        f"de seguros de acuerdo al siguiente detalle:")
    _add_subbullet(doc, f"La Positiva: {fmt_int(d['avisos_lp'])} avisos.")
    _add_subbullet(doc, f"Rímac: {fmt_int(d['avisos_rimac'])} avisos.")

    # Top 4 departamentos
    top4 = d["top4_deptos"]
    top4_items = []
    for depto, count in top4.items():
        pct = (count / d["top4_total"] * 100) if d["top4_total"] > 0 else 0
        top4_items.append(f"{depto.title()} ({fmt_int(count)}, {pct:.1f}%)")

    _add_bullet(doc,
        f"El {d['top4_pct']:.1f}% de avisos reportados se concentra en {len(top4)} "
        f"departamentos: {fmt_int(d['top4_total'])} avisos distribuidos en "
        f"{', '.join(top4_items[:-1])} y {top4_items[-1]}, respectivamente, "
        f"se presenta el detalle de los siniestros según cada departamento:")

    # ─── GRÁFICO 1: Avisos por Departamento ───
    grafico_num += 1
    chart1_buf = _chart_avisos_departamento(d["avisos_by_depto"])
    _add_image_centered(doc, chart1_buf, width_inches=6.2)
    _add_grafico_caption(doc, grafico_num, "Número de Avisos Reportados por Departamento, SAC 2025-2026")

    # Top 3 tipos de siniestro
    top3 = d["top3_tipos"]
    top3_items = []
    for tipo, count in top3.items():
        pct = (count / d["top3_total"] * 100) if d["top3_total"] > 0 else 0
        top3_items.append(f"{tipo.lower()} ({fmt_int(count)}, {pct:.1f}%)")

    _add_bullet(doc,
        f"El {d['top3_pct']:.1f}% de avisos reportados se concentra en {len(top3)} siniestros: "
        f"{fmt_int(d['top3_total'])} avisos distribuidos en "
        f"{', '.join(top3_items[:-1])} y {top3_items[-1]}, "
        f"se presenta el detalle de los avisos reportados según el tipo de siniestro:")

    # ─── GRÁFICO 2: Avisos por Tipo ───
    grafico_num += 1
    chart2_buf = _chart_avisos_tipo(d["avisos_by_tipo"])
    _add_image_centered(doc, chart2_buf, width_inches=6.2)
    _add_grafico_caption(doc, grafico_num, "Número de Avisos Reportados por Tipo de Siniestro, SAC 2025-2026")

    # ─── b) Resultados ───
    _add_section_heading(doc, "b) Resultados")

    _add_bullet(doc,
        f"Del total de avisos de siniestros reportados ({fmt_int(d['total_avisos'])}) a las empresas "
        f"de seguros, al {d['fecha_corte']} se han ajustado y evaluado {fmt_int(d['total_ajustados'])} "
        f"avisos que representa el {d['pct_ajustados']:.2f}% de los avisos.")

    _add_bullet(doc, "Los ajustes y evaluación de los mismos por cada empresa de seguros es la siguiente:")
    _add_subbullet(doc,
        f"La Positiva: {fmt_int(d['ajust_lp'])} ajustes de {fmt_int(d['avisos_lp'])} "
        f"avisos de siniestros, {d['pct_ajust_lp']:.2f}% de avisos atendidos.")
    _add_subbullet(doc,
        f"Rímac: {fmt_int(d['ajust_rimac'])} ajustes de {fmt_int(d['avisos_rimac'])} "
        f"avisos de siniestros, {d['pct_ajust_rimac']:.2f}% de avisos atendidos.")

    # Indemnizaciones (bold + underline)
    _add_bold_underline_bullet(doc,
        f"Las indemnizaciones reconocidas a la fecha, por parte de las empresas de seguros, "
        f"equivalen al monto de {fmt(d['monto_indemnizado'])}.")

    # Siniestralidad
    _add_bullet(doc,
        f"El índice de siniestralidad (costo de los siniestros entre la prima neta, sin incluir IGV, "
        f"de la póliza) en lo que va de la campaña, es de {d['idx_general']:.2f}%, siendo el índice "
        f"para cada empresa de seguros el siguiente:")
    _add_subbullet(doc, f"La Positiva Seguros: {d['idx_lp']:.2f}%.")
    _add_subbullet(doc, f"Rímac Seguros: {d['idx_rimac']:.2f}%.")

    _add_body(doc,
        "Es preciso resaltar que la vigencia de la póliza es desde el 01 de agosto de 2025 "
        "hasta 01 de agosto de 2026, por lo que, todos los eventos adversos que se presenten "
        "durante este periodo, que afecten los cultivos asegurados, sean reportados por las "
        "DRAs/GRAs a las empresas de seguros, para que estas realicen las evaluaciones correspondientes.",
        italic=True, size=9)

    _add_body(doc, "El índice de siniestralidad para cada departamento se detalla en el siguiente cuadro:", size=10)

    # ═══ CUADRO 1: SINIESTRALIDAD ═══
    cuadro_num += 1
    _add_cuadro_title(doc, cuadro_num, "Índice de Siniestralidad por Departamento y Empresa, SAC 2025-2026")

    headers_sin = ["EMPRESA DE SEGUROS", "DEPARTAMENTO", "INDEMNIZACIÓN (S/)",
                   "SUP. INDEMNIZADA (Has)", "PRIMA NETA (S/)", "ÍNDICE SINIESTRALIDAD (%)"]
    col_widths_sin = [1800, 1600, 1700, 1600, 1700, 1500]
    table = _build_table(doc, headers_sin, d["tabla_siniestralidad"], col_widths_sin)

    for row_idx, row_data in enumerate(d["tabla_siniestralidad"]):
        is_sub = row_data.get("is_subtotal", False)
        is_tot = row_data.get("is_total", False)
        bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
        fc = C["WHITE"] if is_tot else C["BLACK"]
        b = is_sub or is_tot
        r = table.rows[row_idx + 1]
        vals = [row_data["empresa"], row_data["departamento"],
                fmt(row_data["indemnizacion"]), fmt_n(row_data["sup_indemnizada"]),
                fmt(row_data["prima_neta"]), fmt_pct(row_data["indice"])]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER]
        for i, (val, align) in enumerate(zip(vals, aligns)):
            _write_cell(r.cells[i], val, bold=b, size=7, align=align, bg_color=bg, font_color=fc)

    doc.add_paragraph()

    # ═══ CUADRO 2: COBERTURAS ═══
    if d["tabla_coberturas"]:
        tc_no_sub = [r for r in d["tabla_coberturas"] if not r.get("is_subtotal") and not r.get("is_total")]
        total_comp = sum(r["complementaria"] for r in tc_no_sub)
        total_cat = sum(r["catastrofica"] for r in tc_no_sub)
        pct_comp = (total_comp / d["monto_indemnizado"] * 100) if d["monto_indemnizado"] > 0 else 0
        pct_cat = (total_cat / d["monto_indemnizado"] * 100) if d["monto_indemnizado"] > 0 else 0

        _add_bullet(doc,
            f"Del total de las indemnizaciones reconocidas por las empresas de seguros, "
            f"se tiene que el {pct_comp:.0f}% es por la cobertura complementaria "
            f"({fmt(total_comp)}); el {pct_cat:.0f}% restante es por la cobertura "
            f"catastrófica de evaluación ({fmt(total_cat)}).")

        cuadro_num += 1
        _add_cuadro_title(doc, cuadro_num, "Indemnizaciones por Tipo de Cobertura, SAC 2025-2026")

        headers_cob = ["EMPRESA DE SEGUROS", "DEPARTAMENTO",
                       "COBERTURA COMPLEMENTARIA", "COBERTURA CATASTRÓFICA", "INDEMNIZACIÓN TOTAL"]
        col_widths_cob = [1800, 1500, 2000, 2000, 1800]
        table_c = _build_table(doc, headers_cob, d["tabla_coberturas"], col_widths_cob)

        for row_idx, row_data in enumerate(d["tabla_coberturas"]):
            is_sub = row_data.get("is_subtotal", False)
            is_tot = row_data.get("is_total", False)
            bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
            fc = C["WHITE"] if is_tot else C["BLACK"]
            b = is_sub or is_tot
            r = table_c.rows[row_idx + 1]
            vals = [row_data["empresa"], row_data.get("departamento", ""),
                    fmt(row_data["complementaria"]), fmt(row_data["catastrofica"]),
                    fmt(row_data["total"])]
            aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                      WD_ALIGN_PARAGRAPH.RIGHT]
            for i, (val, align) in enumerate(zip(vals, aligns)):
                _write_cell(r.cells[i], val, bold=b, size=7, align=align, bg_color=bg, font_color=fc)

        doc.add_paragraph()

    # ═══ CUADRO 3: PRIORIZADOS ═══
    if d["tabla_priorizados"]:
        _add_bullet(doc,
            f"Asimismo, las indemnizaciones que se vienen presentando se puede detallar, "
            f"cultivos priorizados ({fmt(d['total_prio_ind'])}) y cultivos no priorizados "
            f"({fmt(d['total_noprio_ind'])}), las mismas que se detallan de acuerdo al siguiente cuadro:")

        cuadro_num += 1
        _add_cuadro_title(doc, cuadro_num, "Indemnizaciones por Cultivos Priorizados y No Priorizados, SAC 2025-2026")

        headers_p = ["EMPRESA DE SEGUROS", "CULTIVOS", "SUP. INDEMNIZADA (Has)", "INDEMNIZACIÓN (S/)"]
        col_widths_p = [2200, 1800, 2500, 2800]
        table_p = _build_table(doc, headers_p, d["tabla_priorizados"], col_widths_p)

        for row_idx, row_data in enumerate(d["tabla_priorizados"]):
            is_sub = row_data.get("is_subtotal", False)
            is_tot = row_data.get("is_total", False)
            bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
            fc = C["WHITE"] if is_tot else C["BLACK"]
            b = is_sub or is_tot
            r = table_p.rows[row_idx + 1]
            vals = [row_data["empresa"], row_data.get("cultivo", ""),
                    fmt_n(row_data["sup_indemnizada"]), fmt(row_data["indemnizacion"])]
            aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
            for i, (val, align) in enumerate(zip(vals, aligns)):
                _write_cell(r.cells[i], val, bold=b, size=7, align=align, bg_color=bg, font_color=fc)

        doc.add_paragraph()

    # ═══ DESEMBOLSOS + GRÁFICO 3 ═══
    _add_bullet(doc,
        f"Con respecto a las indemnizaciones reconocidas por parte de las empresas de seguros "
        f"del SAC 2025-2026, a la fecha el avance de desembolsos realizados es la siguiente:")

    # ─── GRÁFICO 3: Desembolsos por empresa ───
    grafico_num += 1
    chart3_buf = _chart_desembolsos_empresa(
        d["indemn_lp"], d["desemb_lp"], d["indemn_rimac"], d["desemb_rimac"])
    _add_image_centered(doc, chart3_buf, width_inches=5.5)
    _add_grafico_caption(doc, grafico_num, "Desembolsos Realizados a Nivel de Empresa de Seguros, SAC 2025-2026")

    # Departamentos con desembolso
    deptos_lp = [r["departamento"] for r in d["tabla_desembolsos"]
                 if r.get("empresa") == "LA POSITIVA" and r.get("desembolso", 0) > 0
                 and not r.get("is_subtotal") and not r.get("is_total")]
    deptos_rimac = [r["departamento"] for r in d["tabla_desembolsos"]
                    if r.get("empresa") == "RÍMAC" and r.get("desembolso", 0) > 0
                    and not r.get("is_subtotal") and not r.get("is_total")]

    if deptos_lp:
        _add_body(doc,
            f"A nivel de departamento, La Positiva Seguros ha iniciado con los desembolsos de "
            f"las indemnizaciones en los departamentos de {', '.join(deptos_lp)}.", size=10)
    if deptos_rimac:
        _add_body(doc,
            f"Rímac Seguros ha iniciado con los desembolsos de indemnizaciones en los "
            f"departamentos de {', '.join(deptos_rimac)}.", size=10)

    # ═══ CUADRO 4: DESEMBOLSOS ═══
    cuadro_num += 1
    _add_cuadro_title(doc, cuadro_num, "Desembolsos y Número de Productores, SAC 2025-2026")

    headers_d = ["EMPRESA DE SEGUROS", "DEPARTAMENTO", "INDEMNIZACIÓN (S/)",
                 "DESEMBOLSO (S/)", "% DESEMBOLSO", "N° PRODUCTORES"]
    col_widths_d = [1700, 1500, 1700, 1700, 1200, 1400]
    table_d = _build_table(doc, headers_d, d["tabla_desembolsos"], col_widths_d)

    for row_idx, row_data in enumerate(d["tabla_desembolsos"]):
        is_sub = row_data.get("is_subtotal", False)
        is_tot = row_data.get("is_total", False)
        bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
        fc = C["WHITE"] if is_tot else C["BLACK"]
        b = is_sub or is_tot
        r = table_d.rows[row_idx + 1]
        prod_text = fmt_int(row_data["productores"]) if row_data["productores"] > 0 else "-"
        vals = [row_data["empresa"], row_data.get("departamento", ""),
                fmt(row_data["indemnizacion"]), fmt(row_data["desembolso"]),
                fmt_pct(row_data["pct_desembolso"]), prod_text]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        for i, (val, align) in enumerate(zip(vals, aligns)):
            _write_cell(r.cells[i], val, bold=b, size=7, align=align, bg_color=bg, font_color=fc)

    doc.add_paragraph()

    # ═══ PÁRRAFOS FINALES ═══
    _add_bullet(doc,
        f"A la fecha se van indemnizando a {fmt_int(d['productores'])} productores en "
        f"{d['deptos_con_desembolso']} de los 24 departamentos por el Seguro Agrícola "
        f"Catastrófico en la presente campaña.")

    _add_bullet(doc,
        f"Se va teniendo un porcentaje de desembolso del {float(d['pct_desembolso']):.2f}%.")

    # ═══ GUARDAR ═══
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

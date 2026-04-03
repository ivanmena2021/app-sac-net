"""
gen_word_nacional_py.py — Generador Ayuda Memoria Nacional SAC (python-docx)
==============================================================================
Versión 3.0 — Diseño compacto profesional (máximo 2 páginas)
- Título con franja de color
- Datos generales como tarjeta compacta
- Tablas con fuente 7.5pt para caber en 2 páginas
- Sin sección de activación (se omite para ahorrar espacio)
- Sin filas TOTAL duplicadas
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from io import BytesIO


# ═══════════════════════════════════════════════════════
# PALETA DE COLORES
# ═══════════════════════════════════════════════════════
AZUL = "1F4E79"
AZUL_CLARO = "D6E4F0"
AZUL_MED = "2E75B6"
GRIS = "F2F2F2"
BLANCO = "FFFFFF"
NEGRO = "1A1A1A"
VERDE = "548235"
BORDE = "B4C6D9"

FONT_TABLE = Pt(7.5)
FONT_HEADER = Pt(7.5)
FONT_BODY = Pt(9)
FONT_SMALL = Pt(8)


def fmt(val, dec=2):
    """Formato numérico peruano: 1.234.567,89"""
    if val is None or val == "":
        return "0"
    try:
        n = float(val)
        if dec == 0:
            return f"{n:,.0f}".replace(",", " ").replace(".", ",").replace(" ", ".")
        return f"{n:,.{dec}f}".replace(",", " ").replace(".", ",").replace(" ", ".")
    except (ValueError, TypeError):
        return "0"


def _shading(cell, color):
    """Aplica fondo a celda."""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _borders(cell, color=BORDE, sz="3"):
    """Aplica bordes a celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    xml = (f'<w:tcBorders {nsdecls("w")}>'
           f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'</w:tcBorders>')
    tcPr.append(parse_xml(xml))


def _cell_margins(cell, top=30, bottom=30, left=60, right=60):
    """Márgenes internos de celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def _set_row_height(row, height_twips):
    """Fija altura de fila."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_twips}" w:hRule="atLeast"/>')
    trPr.append(trHeight)


def _add_run(para, text, size=FONT_BODY, bold=False, color=NEGRO, font="Arial Narrow"):
    """Agrega run formateado."""
    r = para.add_run(text)
    r.font.name = font
    r.font.size = size
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _heading(doc, text, level=1):
    """Crea heading compacto con línea azul inferior."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8) if level == 1 else Pt(6)
    p.paragraph_format.space_after = Pt(4)
    size = Pt(11) if level == 1 else Pt(10)
    _add_run(p, text, size=size, bold=True, color=AZUL)
    # Línea inferior
    pPr = p._element.get_or_add_pPr()
    bdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{AZUL_MED}"/></w:pBdr>')
    pPr.append(bdr)
    return p


def _compact_table(doc, headers, rows, col_widths):
    """Crea tabla compacta profesional."""
    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Anchos de columna
    for i, w in enumerate(col_widths):
        table.columns[i].width = Twips(w)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _add_run(p, h, size=FONT_HEADER, bold=True, color=BLANCO)
        _shading(cell, AZUL)
        _borders(cell, AZUL)
        _cell_margins(cell, 25, 25, 40, 40)

    _set_row_height(table.rows[0], 260)

    # Datos
    for r_idx, row_data in enumerate(rows):
        is_total = str(row_data[0]).upper() == "TOTAL"
        is_alt = r_idx % 2 == 0

        table_row = table.rows[r_idx + 1]
        _set_row_height(table_row, 220)

        for c_idx, val in enumerate(row_data):
            cell = table_row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            txt_color = BLANCO if is_total else NEGRO
            _add_run(p, str(val) if val else "", size=FONT_TABLE, bold=is_total, color=txt_color)

            if is_total:
                _shading(cell, AZUL)
            elif is_alt:
                _shading(cell, AZUL_CLARO)

            _borders(cell, BORDE)
            _cell_margins(cell, 20, 20, 40, 40)

    return table


# ═══════════════════════════════════════════════════════
# GENERADOR PRINCIPAL
# ═══════════════════════════════════════════════════════

def generate_nacional_docx(datos):
    """Genera Ayuda Memoria Nacional SAC — diseño compacto 2 páginas."""
    doc = Document()

    # Márgenes estrechos para maximizar espacio
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ═══ TÍTULO ═══
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    _add_run(title_p, "AYUDA MEMORIA", size=Pt(16), bold=True, color=AZUL)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    _add_run(sub_p, "RESUMEN OPERATIVIDAD SAC 2025-2026", size=Pt(12), bold=True, color=AZUL_MED)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(8)
    _add_run(date_p, f"(al {datos['fecha_corte']})", size=Pt(9), color="666666")

    # ═══ DATOS GENERALES — como párrafo denso ═══
    _heading(doc, "Datos Generales a Nivel Nacional")

    # Párrafo consolidado en vez de bullets
    info_p = doc.add_paragraph()
    info_p.paragraph_format.space_after = Pt(4)
    info_p.paragraph_format.line_spacing = Pt(13)

    text = (
        f"Empresas aseguradoras: {datos['empresas_text']}. "
        f"Prima total (con IGV): S/ {fmt(datos['prima_total'])} | "
        f"Prima neta (sin IGV): S/ {fmt(datos['prima_neta'])}. "
        f"Superficie asegurada: {fmt(datos['sup_asegurada'])} ha en 24 departamentos. "
        f"Productores asegurados (est.): {fmt(datos['prod_asegurados'], 0)} | "
        f"Suma asegurada: S/ 1.000/ha."
    )
    _add_run(info_p, text, size=FONT_SMALL, bold=False, color=NEGRO)

    # Línea de indicadores clave
    ind_p = doc.add_paragraph()
    ind_p.paragraph_format.space_after = Pt(6)
    ind_p.paragraph_format.line_spacing = Pt(13)

    ind_text = (
        f"Avisos de siniestros: {fmt(datos['total_avisos'], 0)} reportados | "
        f"{fmt(datos['total_ajustados'], 0)} evaluados ({datos['pct_ajustados']}%). "
        f"Indemnización: S/ {fmt(datos['monto_indemnizado'])} | Siniestralidad: {datos['indice_siniestralidad']}%. "
        f"Desembolsos: S/ {fmt(datos['monto_desembolsado'])} ({datos['pct_desembolso']}%) "
        f"a {fmt(datos['productores_desembolso'], 0)} productores en {datos['deptos_con_desembolso']} de 24 dptos."
    )
    _add_run(ind_p, ind_text, size=FONT_SMALL, bold=True, color=NEGRO)

    # ═══ CUADRO 1: Primas y Cobertura ═══
    _heading(doc, "Cuadro 1: Primas y Cobertura por Departamento")

    if datos.get("cuadro1") and len(datos["cuadro1"]) > 0:
        headers = ["Departamento", "Prima Total (S/)", "Ha Aseguradas", "Suma Asegurada (S/)"]
        rows = []
        for item in datos["cuadro1"]:
            rows.append([
                item.get("departamento", ""),
                fmt(item.get("prima_total", 0)),
                fmt(item.get("hectareas", 0)),
                fmt(item.get("suma_asegurada", 0)),
            ])

        # TOTAL único
        t_prima = sum(float(item.get("prima_total", 0) or 0) for item in datos["cuadro1"])
        t_ha = sum(float(item.get("hectareas", 0) or 0) for item in datos["cuadro1"])
        t_suma = sum(float(item.get("suma_asegurada", 0) or 0) for item in datos["cuadro1"])
        rows.append(["TOTAL", fmt(t_prima), fmt(t_ha), fmt(t_suma)])

        _compact_table(doc, headers, rows, [2100, 1800, 1800, 2100])

    # ═══ CUADRO 2: Indemnizaciones y Desembolsos ═══
    _heading(doc, "Cuadro 2: Indemnizaciones y Desembolsos por Departamento")

    if datos.get("cuadro2") and len(datos["cuadro2"]) > 0:
        headers = ["Departamento", "Ha Indemn.", "Indemnización (S/)", "Desembolso (S/)", "Productores"]
        rows = []
        for item in datos["cuadro2"]:
            rows.append([
                item.get("departamento", ""),
                fmt(item.get("ha_indemnizadas", 0)),
                fmt(item.get("monto_indemnizado", 0)),
                fmt(item.get("monto_desembolsado", 0)),
                fmt(item.get("productores", 0), 0),
            ])

        t_ha = sum(float(item.get("ha_indemnizadas", 0) or 0) for item in datos["cuadro2"])
        t_ind = sum(float(item.get("monto_indemnizado", 0) or 0) for item in datos["cuadro2"])
        t_des = sum(float(item.get("monto_desembolsado", 0) or 0) for item in datos["cuadro2"])
        t_pro = sum(float(item.get("productores", 0) or 0) for item in datos["cuadro2"])
        rows.append(["TOTAL", fmt(t_ha), fmt(t_ind), fmt(t_des), fmt(t_pro, 0)])

        _compact_table(doc, headers, rows, [2100, 1200, 1800, 1800, 900])

    # ═══ CUADRO 3: Lluvias Intensas ═══
    _heading(doc, "Cuadro 3: Eventos Asociados a Lluvias Intensas")

    # Párrafo descriptivo compacto
    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.space_after = Pt(4)
    desc_text = (
        f"Se registran {fmt(datos['total_lluvia'], 0)} avisos por eventos asociados a lluvias intensas "
        f"({datos['pct_lluvia']}% del total), que incluyen {datos['lluvia_desc']}. "
        f"Los departamentos más afectados son {datos['top3_lluvia_text']}."
    )
    _add_run(desc_p, desc_text, size=FONT_SMALL, bold=False)

    if datos.get("cuadro3") and len(datos["cuadro3"]) > 0:
        headers = ["Departamento", "Avisos", "Ha Indemn.", "Indemnización (S/)", "Desembolso (S/)", "Prod."]
        rows = []
        for item in datos["cuadro3"]:
            rows.append([
                item.get("departamento", ""),
                fmt(item.get("avisos", 0), 0),
                fmt(item.get("ha_indemn", 0)),
                fmt(item.get("monto_indemnizado", 0)),
                fmt(item.get("monto_desembolsado", 0)),
                fmt(item.get("productores", 0), 0),
            ])

        _compact_table(doc, headers, rows, [1800, 700, 1000, 1700, 1700, 700])

    # ═══ NOTA FINAL ═══
    doc.add_paragraph()  # pequeño espacio
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_after = Pt(2)
    _add_run(note_p, "Nota: ", size=Pt(7.5), bold=True, color="666666")
    _add_run(note_p, "Vigencia de póliza: 01/08/2025 al 01/08/2026. ", size=Pt(7.5), color="666666")
    _add_run(note_p, datos.get("top3_siniestros_text", ""), size=Pt(7.5), color="666666")

    # Fuente
    src_p = doc.add_paragraph()
    src_p.paragraph_format.space_before = Pt(4)
    _add_run(src_p, "Fuente: Dirección de Seguro y Fomento del Financiamiento Agrario - MIDAGRI, SAC 2025-2026.",
             size=Pt(7), bold=False, color="888888")

    # ═══ RETORNAR BYTES ═══
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

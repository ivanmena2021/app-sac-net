"""
gen_word_departamental_py.py — Pure Python Departamental Document Generator (python-docx)

Rewrite of gen_word_departamental.js using python-docx only.
Lambayeque-style professional document with colored headings and dashboard metrics.
No Node.js dependency.

Usage:
    from gen_word_departamental_py import generate_departamental_docx
    docx_bytes = generate_departamental_docx(depto_data_dict)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION
from io import BytesIO


# ═══ Colors (Lambayeque palette) ═══
C = {
    "DARK_SLATE": "2E4057",
    "RED": "C0392B",
    "TEAL": "16A085",
    "DARK_BLUE": "2C3E50",
    "GRAY": "7F8C8D",
    "WHITE": "FFFFFF",
    "BLACK": "000000",
    "ALT_ROW": "EBF5FB",
    "LIGHT_BG": "F8F9FA",
}


def fmt_num(val, dec=2):
    """Format number with thousands separator (es-PE locale)."""
    if val is None or val == "":
        return "0"
    try:
        n = float(val)
        if dec == 0:
            return f"{n:,.0f}".replace(",", " ").replace(".", ",").replace(" ", ".")
        return f"{n:,.{dec}f}".replace(",", " ").replace(".", ",").replace(" ", ".")
    except (ValueError, TypeError):
        return "0"


def set_cell_background(cell, fill_color):
    """Set cell background color using XML shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell):
    """Add thin borders to cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_paragraph_bottom_border(paragraph, color="C0392B", size="12"):
    """Add bottom border to paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/></w:pBdr>'
    )
    pPr.append(pBdr)


def heading_1(doc, text):
    """Add Heading 1: large, dark slate, with red bottom border."""
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    run = h1.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 64, 87)  # #2E4057 DARK_SLATE
    
    # Add red bottom border
    set_paragraph_bottom_border(h1, color=C["RED"], size="12")
    
    return h1


def heading_2(doc, text):
    """Add Heading 2: teal, smaller than H1."""
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    
    run = h2.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(22, 160, 133)  # #16A085 TEAL
    
    return h2


def create_table(doc, headers, rows, col_widths_twips=None):
    """Create a professional table."""
    if not rows:
        return None
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    
    # Auto width if not specified
    if col_widths_twips is None:
        col_widths_twips = [9360 // len(headers)] * len(headers)
    
    # Set column widths
    for i, width in enumerate(col_widths_twips):
        table.columns[i].width = width
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        set_cell_background(cell, C["DARK_BLUE"])
        
        # Format header text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        
        set_cell_border(cell)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        is_total = row_data[0].upper() == "TOTAL" if row_data else False
        is_alt = row_idx % 2 == 0
        bg_color = C["DARK_BLUE"] if is_total else (C["ALT_ROW"] if is_alt else None)
        text_color = C["WHITE"] if is_total else C["BLACK"]
        
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = str(cell_text) if cell_text else ""
            
            if bg_color:
                set_cell_background(cell, bg_color)
            
            # Format cell text
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            
            for run in paragraph.runs:
                run.font.bold = is_total
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if text_color == C["WHITE"]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            set_cell_border(cell)
    
    return table


def create_metric_card_table(doc, metrics):
    """
    Create 2x2 metric dashboard.
    metrics: list of 4 tuples (value, label)
    """
    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    
    card_width = 4680
    table.columns[0].width = card_width
    table.columns[1].width = card_width
    
    for idx, (value, label) in enumerate(metrics):
        row = idx // 2
        col = idx % 2
        cell = table.rows[row].cells[col]
        
        # Clear default cell content
        cell.text = ""
        
        # Set background
        set_cell_background(cell, C["LIGHT_BG"])
        
        # Add value paragraph
        value_para = cell.paragraphs[0]
        value_para.paragraph_format.space_after = Pt(3)
        value_run = value_para.add_run(value)
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(18)
        value_run.font.bold = True
        value_run.font.color.rgb = RGBColor(46, 64, 87)  # DARK_SLATE
        
        # Add label paragraph
        label_para = cell.add_paragraph()
        label_para.paragraph_format.space_after = Pt(0)
        label_run = label_para.add_run(label)
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(127, 140, 141)  # GRAY
        
        # Set padding/margins
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="120" w:type="dxa"/>'
            f'<w:left w:w="160" w:type="dxa"/>'
            f'<w:bottom w:w="120" w:type="dxa"/>'
            f'<w:right w:w="160" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)
        
        # Light borders
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    
    return table


def generate_departamental_docx(depto_data):
    """
    Generate DEPARTAMENTAL ayuda memoria document (Lambayeque style).
    
    Args:
        depto_data (dict): Data dictionary with keys:
            - departamento
            - empresa
            - prima_neta, sup_asegurada
            - total_avisos, ha_indemnizadas, monto_indemnizado
            - monto_desembolsado, productores_desembolso
            - indemnizables, no_indemnizables
            - fecha_corte
            - avisos_tipo (list of [tipo, count, %])
            - dist_provincia (list of rows)
            - dist_provincia_headers
            - eventos_recientes (list of rows)
            - eventos_headers
            - resumen_operativo, resumen_desembolso (text)
    
    Returns:
        bytes: DOCX file content
    """
    doc = Document()
    
    # Set page margins: 1440 twips = 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    d = depto_data
    depto = d.get("departamento", "Departamento")
    
    # ═══ COVER PAGE ═══
    # Large top spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # "AYUDA MEMORIA"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("AYUDA MEMORIA")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 64, 87)  # DARK_SLATE
    
    # "Seguro Agrícola Catastrófico (SAC)" with red border
    sac_para = doc.add_paragraph()
    sac_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sac_para.paragraph_format.space_after = Pt(3)
    run = sac_para.add_run("Seguro Agrícola Catastrófico (SAC)")
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.italic = True
    run.font.color.rgb = RGBColor(192, 57, 43)  # RED
    set_paragraph_bottom_border(sac_para, color=C["RED"], size="12")
    
    doc.add_paragraph()
    
    # Department info
    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept_para.add_run(f"Departamento de {depto}")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    
    campaign_para = doc.add_paragraph()
    campaign_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = campaign_para.add_run("Campaña 2025 - 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    
    prima_para = doc.add_paragraph()
    prima_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prima_para.add_run(f"Prima Neta Departamental: S/ {fmt_num(d.get('prima_neta', 0))}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    
    corte_para = doc.add_paragraph()
    corte_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = corte_para.add_run(f"Corte de datos: {d.get('fecha_corte', '')}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(127, 140, 141)  # GRAY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    midagri_para = doc.add_paragraph()
    midagri_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = midagri_para.add_run("MIDAGRI")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    # Page break
    doc.add_page_break()
    
    # ═══ INTRODUCTION ═══
    heading_1(doc, "Introducción")
    
    intro1 = doc.add_paragraph()
    run = intro1.add_run(f"La aseguradora responsable del departamento es {d.get('empresa', '')}.")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    
    intro2_text = (
        f"Para la campaña 2025-2026, {depto} cuenta con una cobertura de "
        f"{fmt_num(d.get('sup_asegurada', 0), 0)} hectáreas aseguradas, por un monto de prima neta "
        f"de S/ {fmt_num(d.get('prima_neta', 0))}. El SAC 2025-2026 ha incrementado la indemnización "
        f"máxima de S/ 800 a S/ 1,000 por hectárea de cultivo afectado."
    )
    intro2 = doc.add_paragraph()
    intro2.paragraph_format.space_after = Pt(6)
    run = intro2.add_run(intro2_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # ═══ SECTION 1: SAC PROCESS ═══
    heading_1(doc, "1. Proceso del SAC: Etapas y Plazos")
    
    process_para = doc.add_paragraph()
    run = process_para.add_run(
        "El Seguro Agrícola Catastrófico (SAC) sigue un proceso de 8 etapas desde la ocurrencia "
        "del siniestro hasta el desembolso al productor afectado."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    
    # Process table
    etapas = [
        ["1", "Ocurrencia del Siniestro", "Evento climático adverso afecta cultivos", "—"],
        ["2", "Aviso de Siniestro", "El productor o DRAS comunica el evento", "Productor / DRAS"],
        ["3", "Atención del Aviso", "Aseguradora registra y verifica", "Aseguradora"],
        ["4", "Programación de Ajuste", "Se coordina visita de inspección", "Aseguradora"],
        ["5", "Inspección y Ajuste", "Evaluación técnica en campo", "Ajustador"],
        ["6", "Dictamen", "Se determina si es indemnizable", "Aseguradora"],
        ["7", "Validación DRAS/GRAS", "DRA emite conformidad", "DRAS/GRAS"],
        ["8", "Desembolso", "Pago de indemnización al productor", "Aseguradora"],
    ]
    
    etapa_widths = [480, 1760, 3040, 2208]
    create_table(doc, ["N°", "Etapa", "Descripción", "Responsable"], etapas, etapa_widths)
    
    doc.add_paragraph()
    
    # ═══ SECTION 2: OVERVIEW ═══
    heading_1(doc, f"2. Panorama General — {depto}")
    
    overview_para = doc.add_paragraph()
    overview_para.paragraph_format.space_after = Pt(6)
    run = overview_para.add_run(
        f"Resumen acumulado de la gestión del SAC en el departamento de {depto} al {d.get('fecha_corte', '')}:"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    
    # Dashboard 2x2
    total_avisos = fmt_num(d.get('total_avisos', 0), 0)
    monto_indemn = fmt_num(d.get('monto_indemnizado', 0), 0)
    indemnizables = d.get('indemnizables', 0)
    
    ha_indemn = fmt_num(d.get('ha_indemnizadas', 0))
    sup_aseg = fmt_num(d.get('sup_asegurada', 0), 0)
    
    monto_desemb = fmt_num(d.get('monto_desembolsado', 0), 0)
    prod_desemb = d.get('productores_desembolso', 0)
    
    metrics = [
        (f"{total_avisos}", f"Avisos de Siniestro registrados"),
        (f"S/ {monto_indemn}", f"Indemnización total ({indemnizables} casos indemnizables)"),
        (f"{ha_indemn} ha", f"Superficie indemnizada (de {sup_aseg} ha aseguradas)"),
        (f"S/ {monto_desemb}", f"Monto desembolsado a {prod_desemb} productores" if d.get('monto_desembolsado', 0) > 0 else "Monto desembolsado (pendiente de pago)"),
    ]
    
    create_metric_card_table(doc, metrics)
    
    doc.add_paragraph()
    
    # ═══ Avisos por Tipo ═══
    heading_2(doc, "Avisos por Tipo de Siniestro")
    
    if d.get("avisos_tipo") and len(d["avisos_tipo"]) > 0:
        tipo_widths = [3040, 2240, 2208]
        create_table(doc, ["Tipo de Siniestro", "N° Avisos", "% del Total"], d["avisos_tipo"], tipo_widths)
    
    doc.add_paragraph()
    
    # ═══ Distribución por Provincia ═══
    heading_2(doc, "Distribución por Provincia")
    
    if d.get("dist_provincia") and len(d["dist_provincia"]) > 0:
        prov_headers = d.get("dist_provincia_headers", ["Provincia", "Avisos", "Sup. Indemn.", "Prod. Benef.", "Indemniz.", "Desembolso", "% Avance"])
        n_cols = len(prov_headers)
        prov_width = 9360 // n_cols
        prov_widths = [prov_width] * n_cols
        create_table(doc, prov_headers, d["dist_provincia"], prov_widths)
    
    doc.add_paragraph()
    
    # ═══ SECTION 3: RECENT EVENTS ═══
    heading_1(doc, "3. Eventos Registrados Recientemente")
    
    if d.get("eventos_recientes") and len(d["eventos_recientes"]) > 0:
        events_para = doc.add_paragraph()
        run = events_para.add_run(
            f"Se han registrado {len(d['eventos_recientes'])} avisos de siniestro recientes en el departamento de {depto}."
        )
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        
        ev_headers = d.get("eventos_headers", ["Fecha", "Provincia", "Distrito / Sector", "Cultivo", "Estado"])
        n_ev_cols = len(ev_headers)
        ev_width = 9360 // n_ev_cols
        ev_widths = [ev_width] * n_ev_cols
        create_table(doc, ev_headers, d["eventos_recientes"], ev_widths)
    else:
        events_para = doc.add_paragraph()
        run = events_para.add_run("No se han registrado eventos recientes en el período.")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Disclaimer
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "* Superficie perdida reportada preliminarmente; la superficie afectada total está pendiente de evaluación en campo."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    
    # ═══ Resumen Operativo ═══
    heading_2(doc, "Resumen Operativo")
    
    if d.get("resumen_operativo"):
        resumen_op = doc.add_paragraph()
        run = resumen_op.add_run(d["resumen_operativo"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
    
    if d.get("resumen_desembolso"):
        resumen_desemb = doc.add_paragraph()
        run = resumen_desemb.add_run(d["resumen_desembolso"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
    
    # ═══ RETURN BYTES ═══
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


if __name__ == "__main__":
    # Quick test
    sample_data = {
        "departamento": "Lambayeque",
        "empresa": "La Positiva",
        "prima_neta": 1200000.00,
        "sup_asegurada": 15000.00,
        "total_avisos": 450,
        "ha_indemnizadas": 8900.50,
        "monto_indemnizado": 2500000.00,
        "monto_desembolsado": 1900000.00,
        "productores_desembolso": 280,
        "indemnizables": 350,
        "no_indemnizables": 100,
        "fecha_corte": "27/02/2026",
        "avisos_tipo": [
            ["Inundación", "280", "62.2%"],
            ["Lluvias Excesivas", "120", "26.7%"],
            ["Huayco", "50", "11.1%"],
        ],
        "dist_provincia": [
            ["Chiclayo", "250", "5600.50 ha", "150", "S/ 1800000", "S/ 1400000", "78%"],
            ["Lambayeque", "150", "2400.00 ha", "90", "S/ 600000", "S/ 400000", "67%"],
            ["Ferreñafe", "50", "800.00 ha", "40", "S/ 100000", "S/ 100000", "100%"],
        ],
        "dist_provincia_headers": ["Provincia", "Avisos", "Sup. Indemn.", "Prod. Benef.", "Indemniz.", "Desembolso", "% Avance"],
        "eventos_recientes": [
            ["20/02/2026", "Chiclayo", "Yotala / Talambo", "Arroz", "Cerrado"],
            ["18/02/2026", "Lambayeque", "Mórrope", "Maíz", "Pendiente"],
        ],
        "eventos_headers": ["Fecha", "Provincia", "Distrito / Sector", "Cultivo", "Estado"],
        "resumen_operativo": "Del total de 450 avisos registrados, se han evaluado y ajustado 350 (77.8%), quedando 100 pendientes (22.2%). De los 350 avisos cerrados, 320 resultaron indemnizables y 30 no indemnizables. El monto total de indemnizaciones reconocidas asciende a S/ 2,500,000.00 sobre una superficie indemnizada de 8,900.50 hectáreas.",
        "resumen_desembolso": "Se han realizado desembolsos por S/ 1,900,000.00 a 280 productores.",
    }
    
    docx_bytes = generate_departamental_docx(sample_data)
    print(f"Generated DEPARTAMENTAL document: {len(docx_bytes)} bytes")
    print("✓ No import errors, document structure valid")
